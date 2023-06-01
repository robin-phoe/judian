from typing import List

import win32api
import win32print
import datetime

from docx import Document
from docx.shared import Cm,Pt
import pathlib
import os

PRROJET_DIR = pathlib.Path(__file__).resolve().parent.parent.parent.parent
BILL_FILE_NAME = os.path.join(PRROJET_DIR,'Judian','bill.docx')
def do_print():

    win32api.ShellExecute(
        0,
        "print",
        BILL_FILE_NAME,
        #
        # If this is None, the default printer will
        # be used anyway.
        # 使用默认打印机
        '/d:"%s"' % win32print.GetDefaultPrinter(),
        ".",
        0
    )


#订单内容格式化
def print_bill(order_id, time_str, order_infos:List, total_price):
    NameLenth = 10
    CountLenth = 5
    PriceLenth = 6
    TotalPriceLenth = 8
    order_content = ""
    for order in order_infos:
        order_content += order['name'].ljust(NameLenth - len(order['name']), ' ')
        order_content += str(order['count']).center(CountLenth, ' ')
        order_content += str(order['price']).center(PriceLenth, ' ')
        order_content += str(order['total_price']).rjust(TotalPriceLenth, ' ')
        order_content += '\n'
    print(order_content)
    edit_print_doc(order_id, time_str, order_content, total_price)
    do_print()
#编辑打印文档文件
def edit_print_doc(order_id, time_str, order_content, total_price):
    doc = Document()
    p1 = doc.add_paragraph()
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5

    t1 = p1.add_run('     聚点棋牌')
    t1.font.name = '宋体'
    t1.font.size = Pt(14)
    t1.blod = True

    t1 = p1.add_run('*******************\n')
    t1.font.name = '宋体'
    t1.font.size = Pt(14)
    t1.blod = True

    t1 = p1.add_run('-----------------------------\n')
    t1.font.name = '宋体'
    t1.font.size = Pt(9)
    t1.blod = True

    t1 = p1.add_run('单号:{}\n'.format(order_id))
    t1.font.name = '宋体'
    t1.font.size = Pt(9)
    t1.blod = True

    t1 = p1.add_run('时间:{}\n'.format(time_str))
    t1.font.name = '宋体'
    t1.font.size = Pt(9)
    t1.blod = True

    t1 = p1.add_run('-----------------------------\n')
    t1.font.name = '宋体'
    t1.font.size = Pt(9)
    t1.blod = True

    t1 = p1.add_run('名称      数量   单价   金额\n')
    t1.font.name = '宋体'
    t1.font.size = Pt(9)
    t1.blod = True

    # t1 = p1.add_run('哇哈哈      1     5       5\n套餐        1    100     100\n')
    t1 = p1.add_run('{}'.format(order_content))
    t1.font.name = '宋体'
    t1.font.size = Pt(9)

    t1 = p1.add_run('-----------------------------\n')
    t1.font.name = '宋体'
    t1.font.size = Pt(9)
    t1.blod = True

    t1 = p1.add_run('                总额:{}\n'.format(total_price))
    t1.font.name = '宋体'
    t1.font.size = Pt(9)
    t1.blod = True

    t1 = p1.add_run('*******************\n')
    t1.font.name = '宋体'
    t1.font.size = Pt(14)
    t1.blod = True

    t1 = p1.add_run('   谢谢惠顾，欢迎下次光临！\n')
    t1.font.name = '宋体'
    t1.font.size = Pt(9)
    t1.blod = True

    default_section = doc.sections[0]
    # 默认宽度和高度
    print(default_section.page_width.cm)  # 21.59
    print(default_section.page_height.cm)  # 27.94
    # 可直接修改宽度和高度，即纸张大小改为自定义
    default_section.page_width = Cm(5.7)
    default_section.page_height = Cm(20)

    print(default_section.top_margin.cm)  # 2.54
    print(default_section.right_margin.cm)  # 3.175
    print(default_section.bottom_margin.cm)  # 2.54
    print(default_section.left_margin.cm)  # 3.175
    # 修改页边距
    default_section.top_margin = Cm(0.5)
    default_section.right_margin = Cm(0.5)
    default_section.bottom_margin = Cm(0.5)
    default_section.left_margin = Cm(0.5)

    doc.save(BILL_FILE_NAME)

if __name__ == '__main__':
    print_bill(order_id='000000000001', time_str=datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                         order_infos=[{"name":"哇哈哈哈哈","count":'5',"price":'10',"total_price":'50.00'},{"name":"红牛","count":'5',"price":'10',"total_price":'150.00'}],
                         total_price = "100.00")
    # print(PRROJET_DIR)